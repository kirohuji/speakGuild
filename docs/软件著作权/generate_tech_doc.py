"""漫语町（ManYu）— 技术设计说明书（软件著作权，~50页，正文12pt，1.5倍行距）"""
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

OUTPUT = r'C:\Users\z1309\Desktop\work\speakGuild\docs\软件著作权\ManYu技术文档.docx'
DIAGRAMS = r'C:\Users\z1309\Desktop\work\speakGuild\docs\软件著作权\diagrams'
HEADER = '漫语町（ManYu）多语种导游资格面试练习系统 V1.0  技术文档'
SNAME = '漫语町（ManYu）多语种导游资格面试练习系统 V1.0'

def hdr(s, pg):
    h = s.header; h.is_linked_to_previous = False
    p = h.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f'{HEADER}  第 {pg} 页')
    r.font.name = 'SimSun'; r.font.size = Pt(9); r.font.color.rgb = RGBColor(128, 128, 128)
    r.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

def ftr(s):
    f = s.footer; f.is_linked_to_previous = False
    p = f.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(SNAME)
    r.font.name = 'SimSun'; r.font.size = Pt(8); r.font.color.rgb = RGBColor(128, 128, 128)
    r.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

def marg(s):
    s.page_width, s.page_height = Cm(21.0), Cm(29.7)
    s.top_margin = Cm(2.54); s.bottom_margin = Cm(2.54)
    s.left_margin = Cm(3.18); s.right_margin = Cm(3.18)

def np(doc, pg):
    s = doc.add_section(); marg(s); hdr(s, pg); ftr(s); return s

def H(doc, t, lv=1):
    h = doc.add_heading(t, level=lv)
    for r in h.runs: r.font.color.rgb = RGBColor(0, 0, 0)

def P(doc, t, b=False, sz=None, indent=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.5
    if indent: p.paragraph_format.first_line_indent = Cm(0.75)
    r = p.add_run(t); r.bold = b
    r.font.name = 'SimSun'; r.font.size = sz or Pt(12)
    r.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    return p

def B(doc, t):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(2); p.paragraph_format.line_spacing = 1.5
    p.clear()
    r = p.add_run(t); r.font.name = 'SimSun'; r.font.size = Pt(12)
    r.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

def T(doc, hds, rows):
    t = doc.add_table(rows=1, cols=len(hds))
    t.style = 'Light Grid Accent 1'; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(hds):
        c = t.rows[0].cells[i]; c.text = ''
        r = c.paragraphs[0].add_run(h); r.bold = True
        r.font.name = 'SimSun'; r.font.size = Pt(10)
    for rd in rows:
        row = t.add_row()
        for i, v in enumerate(rd):
            c = row.cells[i]; c.text = ''
            r = c.paragraphs[0].add_run(str(v))
            r.font.name = 'SimSun'; r.font.size = Pt(10)
    doc.add_paragraph()

def diagram(doc, filename, caption):
    path = os.path.join(DIAGRAMS, filename)
    if os.path.exists(path):
        doc.add_picture(path, width=Inches(5.2))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.0
    r = p.add_run(caption); r.bold = True; r.font.name = 'SimSun'; r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(80, 80, 80); r.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

def main():
    print('Generating technical document...')
    doc = Document(); pg = 1
    s = doc.sections[0]; marg(s); hdr(s, 0); ftr(s)

    # ── Cover page ──
    for _ in range(5): doc.add_paragraph()
    for t, sz, b in [('软件技术设计说明书', 26, True), (SNAME, 16, False), ('版本：V1.0', 14, False), ('编制日期：2026年5月27日', 14, False)]:
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(t); r.font.name = 'SimSun'; r.font.size = Pt(sz); r.bold = b
        r.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体'); doc.add_paragraph()
    doc.add_page_break()

    # ── TOC ──
    np(doc, pg); pg += 1
    P(doc, '目  录', b=True, sz=Pt(16))
    doc.add_paragraph()
    toc = [
        '1  引言','1.1  编写目的','1.2  项目背景','1.3  适用范围',
        '2  系统概述','2.1  系统定位','2.2  核心功能','2.3  用户角色',
        '3  总体架构设计','3.1  架构设计原则','3.2  系统架构图','3.3  四层架构说明','3.4  前后端通信机制',
        '4  技术选型','4.1  前端技术栈','4.2  后端技术栈','4.3  数据库技术','4.4  第三方服务',
        '5  模块详细设计','5.1  模块划分总览','5.2  模块关系图','5.3  认证模块','5.4  练习模块',
        '5.5  场景化学习模块','5.6  AI反馈模块','5.7  TTS语音模块','5.8  互动剧本模块',
        '5.9  模考模块','5.10  会员支付模块','5.11  通知与反馈模块','5.12  管理后台模块',
        '6  数据库设计','6.1  数据模型总览','6.2  核心表设计','6.3  ER关系图','6.4  索引设计','6.5  数据一致性',
        '7  API接口设计','7.1  设计规范','7.2  请求处理流程','7.3  统一响应格式','7.4  SSE流式响应',
        '8  安全设计','8.1  认证流程','8.2  权限控制','8.3  数据安全','8.4  隐私合规',
        '9  核心业务流程','9.1  场景学习流程','9.2  用户练题流程','9.3  模拟考试流程','9.4  TTS语音流程','9.5  AI反馈时序',
        '9.6  支付开通流程','9.7  互动剧本流程',
        '10  部署与运维','10.1  CI/CD流程','10.2  容器化方案','10.3  部署架构',
        '11  技术难点与解决方案','11.1  TTS多厂商抽象','11.2  AI流式反馈','11.3  Ink剧本引擎',
        '11.4  Monorepo管理','11.5  语块三阶段模型','11.6  文件去重',
    ]
    for x in toc:
        indent = len(x) - len(x.lstrip())
        P(doc, '  ' * indent + x)

    # ═══════ 1 引言 ═══════
    np(doc, pg); pg += 1; H(doc, '1  引言')
    H(doc, '1.1  编写目的')
    P(doc, '本文档是漫语町（ManYu）多语种导游资格面试练习系统的技术设计说明书。编写目的是为了全面、系统、准确地描述该软件的技术架构、模块组成、数据库设计、接口规范、安全机制以及核心业务流程，为软件的开发、测试、部署、运维提供技术依据，同时作为计算机软件著作权登记申请的技术证明材料。')
    P(doc, '本文档面向的读者包括：系统设计人员、开发工程师、测试工程师、运维人员以及软件著作权登记审查人员。')
    H(doc, '1.2  项目背景')
    P(doc, '随着我国入境旅游市场的持续增长，多语种导游人才需求日益旺盛。全国导游资格外语面试是多语种导游获取执业资格的关键环节。考试采用人机对话或考官面试形式，考查内容包括景点讲解、导游服务规范、应变能力、综合知识和外语口译五个维度。目前全国31个省市自治区的考试大纲各有差异，备考资源分散，考生难以高效备考。')
    P(doc, '基于上述背景，本项目致力于打造一款面向外语导游资格面试的全栈数字化备考训练系统，帮助考生突破地域限制，高效、精准、个性化地备考。')
    H(doc, '1.3  适用范围')
    P(doc, '本系统适用于参加全国导游资格外语面试（现场考试）的考生，覆盖英语、日语、俄语、法语、德语、西班牙语、朝鲜语、泰语等语种。系统支持Web浏览器和iOS移动端两种使用方式，满足考生在不同场景下的备考需求。')

    # ═══════ 2 系统概述 ═══════
    np(doc, pg); pg += 1; H(doc, '2  系统概述')
    H(doc, '2.1  系统定位')
    P(doc, '漫语町（ManYu）是一款面向导游资格外语面试考生的沉浸式在线练习与备考平台。系统以"场景驱动、语块积累、互动输出"为核心学习理念，围绕场景学习、口语练习、互动剧本、AI反馈的学习闭环，提供题库浏览、场景化学习、口语练习、模拟考试、AI发音反馈、TTS语音合成、互动剧本探险等完整功能链。系统采用SaaS模式运营，支持会员订阅制，管理员可通过Web后台进行内容管理和运营配置。')
    H(doc, '2.2  核心功能')
    P(doc, '题库浏览功能：用户可按省份、语种、考试类型、面试形式等多维度筛选题目，支持关键词搜索，快速定位目标练习内容。题库内容以卡片形式展示标题、难度、时长、摘要和掌握度进度条。')
    P(doc, '场景化学习功能：以"场景（Scene）"为学习单元，每个场景包含场景词汇、语块和训练主题。语块采用三阶段掌握度模型（activated→can_read→can_output），配合例句、翻译和语义上下文帮助深度理解。')
    P(doc, '口语练习功能：逐题练习模式，支持学习模式和练习模式两种训练方式。集成浏览器TTS语音播放、单词点击释义、中英文翻译显隐切换等功能，支持键盘快捷键操作。')
    P(doc, '互动剧本功能：基于Ink脚本引擎的互动式剧情学习，在不同场景地图和地点中与NPC角色展开对话练习。支持自由对话模式和剧情驱动模式，对话记录自动保存并可回溯复习。')
    P(doc, '模拟考试功能：标准化试卷模考，支持标准卷和强化卷两种难度等级。考试结束后自动评分，生成各题型得分明细和薄弱环节分析报告，记录历史成绩趋势。')
    P(doc, 'AI发音反馈功能：集成DeepSeek大语言模型，对用户的口语回答进行多维度实时评分，采用SSE流式输出技术逐字显示评分结果。同时支持AI教学和词汇拓展。')
    P(doc, '会员系统功能：支持标准会员和进阶会员两级套餐，按月/年订阅。不同等级享受不同次数的AI点评、录音上传、错题分析等权益。')
    P(doc, '管理后台功能：提供用户管理、会员管理、账单管理、题库内容管理、资料库管理、消息通知、反馈管理、数据统计看板和系统设置等功能，支撑日常运营维护。')
    H(doc, '2.3  用户角色')
    T(doc, ['角色','权限','功能范围'],[
        ('普通用户','题库浏览、场景学习、口语练习、互动剧本','前端全部页面'),
        ('会员用户','普通用户权限 + AI反馈、词汇拓展、无限次练习','前端全部页面 + AI功能'),
        ('管理员','全部用户权限 + 后台管理','前端页面 + 管理后台'),
    ])

    # ═══════ 3 总体架构设计 ═══════
    np(doc, pg); pg += 1; H(doc, '3  总体架构设计')
    H(doc, '3.1  架构设计原则')
    P(doc, '前后端分离原则：前端仅负责用户界面渲染和交互响应，所有业务逻辑和数据操作均收敛于后端API服务。前端通过HTTP协议与后端通信，不直接访问数据库。这种分离使得前端和后端可以独立开发、测试和部署，互不影响。')
    P(doc, '模块化组织原则：后端采用NestJS的模块化架构，每个业务域独立为一个Module。模块内部遵循Controller（路由控制）→ Service（业务逻辑）→ DTO（数据校验）的三层结构。模块之间通过依赖注入建立关联，避免循环依赖。')
    P(doc, '接口隔离原则：业务模块之间通过抽象接口依赖而非具体实现。例如TTS模块依赖抽象的TtsProvider接口，各厂商分别实现该接口，新增厂商时无需修改调用方代码。')
    H(doc, '3.2  系统架构图')
    diagram(doc, '01-architecture.png', '图1  系统总体架构图')
    P(doc, '系统采用经典的四层浏览器/服务器架构。客户端层包括Web浏览器和iOS原生应用；接入层使用Nginx反向代理统一入口；服务层承载全部业务逻辑；数据层包括PostgreSQL数据库和腾讯云COS对象存储。各层职责清晰、松耦合设计。')
    H(doc, '3.3  四层架构说明')
    P(doc, '客户端层：包含Web浏览器端（基于React 19构建的单页面应用）和iOS端（通过Capacitor框架打包的原生应用）。前端适配桌面端和移动端两种分辨率，统一使用Hash路由模式，确保静态部署和iOS场景均可正常工作。')
    P(doc, '接入层：Nginx作为反向代理服务器，统一接收所有客户端请求。配置/api路径转发至后端API服务，其他路径直接返回前端静态资源。静态资源通过CDN加速分发，提升全球访问速度。')
    P(doc, '服务层：基于NestJS 11构建的RESTful API服务，共25个业务模块，分为核心学习、游戏化引擎、运营支撑、管理&公共四大类别。使用Better Auth框架管理用户认证与会话，Prisma 6作为数据访问层。AI反馈接口采用SSE流式协议，Ink互动剧本引擎驱动剧情分支学习。')
    P(doc, '数据层：PostgreSQL 16作为主数据库，通过Prisma ORM进行类型安全的数据操作。35+个数据模型覆盖全部业务域，支持JSON字段存储半结构化数据（如用户偏好、AI反馈结果、语块掌握度），GIN索引加速数组字段查询。文件存储使用腾讯云COS，通过STS临时密钥实现前端直传。')
    H(doc, '3.4  前后端通信机制')
    P(doc, '前端通过Axios HTTP客户端与后端通信。请求拦截器自动从localStorage读取Bearer Token并注入请求头。响应拦截器自动解包统一响应体{code, message, data}，在遇到401状态码时自动跳转至登录页。AI反馈等特殊接口使用fetch API的ReadableStream接口接收SSE流式数据，逐段解析并实时渲染。')

    # ═══════ 4 技术选型 ═══════
    np(doc, pg); pg += 1; H(doc, '4  技术选型')
    H(doc, '4.1  前端技术栈')
    T(doc, ['技术','版本','用途'],[
        ('React','19','UI组件库，函数式组件 + Hooks'),
        ('TypeScript','5.9','类型安全，开发效率提升'),
        ('Vite','6','构建工具，ESBuild预编译，HMR热更新'),
        ('shadcn/ui','—','UI组件库，基于Radix UI + Tailwind CSS'),
        ('Tailwind CSS','4','原子化CSS框架，语义化设计变量'),
        ('Zustand','5','全局状态管理，轻量按需订阅'),
        ('React Router','7','前端路由，Hash模式兼容iOS'),
        ('Recharts','2','数据可视化图表库'),
        ('Axios','1','HTTP请求封装，拦截器机制'),
    ])
    H(doc, '4.2  后端技术栈')
    T(doc, ['技术','版本','用途'],[
        ('NestJS','11','Node.js后端框架，模块化 + DI'),
        ('TypeScript','5.9','全栈TypeScript，前后端类型共享'),
        ('Prisma','6','ORM框架，Schema声明式建模'),
        ('PostgreSQL','16','关系型数据库，JSON + GIN索引'),
        ('Better Auth','1.6','认证框架，Session + RBAC'),
        ('DeepSeek API','—','大语言模型，口语评分与反馈'),
        ('MiniMax TTS','—','语音合成服务，多音色'),
        ('Cartesia TTS','—','语音合成服务（降级备用）'),
        ('腾讯云COS','—','对象存储与CDN加速'),
    ])
    H(doc, '4.3  数据库技术')
    P(doc, '选择PostgreSQL 16的原因：对JSON数据类型的原生支持，方便存储用户偏好设置、AI反馈结果、语块掌握度等结构化数据；完善的ACID事务支持，保障支付、会员开通等关键操作的原子性；GIN索引加速数组类型字段的包含查询（如QuestionItem的keywords、usedSceneIds等字段）；丰富的扩展生态支持全文检索、地理空间等高级功能。')
    H(doc, '4.4  第三方服务')
    T(doc, ['服务商','用途','对接方式'],[
        ('DeepSeek','AI口语评分','REST API + SSE流式'),
        ('MiniMax','TTS语音合成','REST API + HTTP回调'),
        ('Cartesia','TTS降级备用','WebSocket + REST API'),
        ('腾讯云COS','文件存储','SDK + STS临时密钥'),
        ('支付宝','在线支付','SDK + 异步回调通知'),
    ])

    # ═══════ 5 模块详细设计 ═══════
    np(doc, pg); pg += 1; H(doc, '5  模块详细设计')
    H(doc, '5.1  模块划分总览')
    P(doc, '后端系统共由25个NestJS业务模块组成，按职责划分为核心学习、游戏化引擎、运营支撑、管理&公共四大类别。以下为各类别包含的模块清单。')
    T(doc, ['类别','模块','职责'],[
        ('核心学习','Auth','用户注册登录、Session管理、角色控制'),
        ('核心学习','Practice','练习记录、进度追踪、收藏生词'),
        ('核心学习','PracticeAI','DeepSeek AI口语评分 (SSE流式)'),
        ('核心学习','TTS','多厂商语音合成、配置哈希缓存'),
        ('核心学习','MockExam','模考试卷管理、自动评分、薄弱分析'),
        ('核心学习','QuestionBank','题库、专题、题目四层树形结构'),
        ('核心学习','Chunk','语块三阶段掌握度管理'),
        ('游戏化引擎','Scene','场景分类、前置条件、详情聚合'),
        ('游戏化引擎','Script','剧本章节、Ink脚本解析、核心词汇'),
        ('游戏化引擎','Exploration','地图地点、NPC对话、探险记录'),
        ('游戏化引擎','Learning','学习路径、推荐引擎'),
        ('游戏化引擎','Level','用户等级、经验值计算'),
        ('游戏化引擎','Expression','表达库管理'),
        ('运营支撑','Membership','会员套餐定义、用户会员状态管理'),
        ('运营支撑','Pay','订单创建、支付回调、状态同步'),
        ('运营支撑','Coupon','优惠券创建、使用核销'),
        ('运营支撑','Referral','邀请码、邀请记录、奖励发放'),
        ('运营支撑','Notification','站内通知广播/定向发送'),
        ('运营支撑','Feedback','用户反馈提交与处理'),
        ('管理&公共','Admin','用户/会员/账单/内容/系统设置'),
        ('管理&公共','FileAssets','SHA256去重、引用计数文件管理'),
        ('管理&公共','ResourceLibrary','树形资料库、多文件类型支持'),
        ('管理&公共','Achievement','成就里程碑定义与自动解锁'),
        ('管理&公共','Leaderboard','学习排行榜聚合'),
        ('管理&公共','ConfigGuide','题库绑定配置管理'),
    ])
    H(doc, '5.2  模块关系图')
    diagram(doc, '03-module-structure.png', '图2  后端模块结构图')
    P(doc, '上图展示了25个模块的类别归属和主要依赖关系。核心学习模块依赖认证模块提供的用户身份信息；游戏化引擎模块以Scene模块为中心，关联Script、Exploration、Chunk等模块形成沉浸式学习闭环；运营模块依赖支付网关回调；管理模块依赖文件资产管理。模块间通过NestJS的依赖注入（DI）和模块导入（imports）机制建立关联，避免循环依赖。')
    H(doc, '5.3  认证模块')
    P(doc, '认证模块基于Better Auth框架实现。提供邮箱密码注册登录、OTP邮箱验证、Session令牌管理、用户角色控制等基础能力。所有业务接口通过requireAuthSession()中间件统一鉴权，该函数从请求头中提取Bearer Token，解码后获取用户ID和角色信息，注入Request对象供后续处理使用。管理员接口额外校验用户角色是否为admin。')
    H(doc, '5.4  练习模块')
    P(doc, '练习模块管理用户的逐题练习行为。每次练习生成一条PracticeRecord记录，包含用户ID、题目ID、操作类型（listen/speak/answer）和练习时间。练习进度由PracticeProgress表追踪，记录每道题的掌握度分数（masteryScore，0-100）和最近练习时间。支持收藏题目和生词本功能。')
    H(doc, '5.5  场景化学习模块')
    P(doc, '场景化学习模块以"场景（Scene）"为基本学习单元组织学习内容。每个场景包含三部分内容：场景词汇（SceneVocabulary）、语块（Chunk）和训练主题（TrainingTopic）。场景设有前置场景要求（prerequisiteScenes）和等级要求（requiredUserLevel），形成渐进式的学习路径。用户场景进度（UserSceneProgress）追踪每个场景的完成度和掌握率。')
    P(doc, '语块模块（ChunkModule）管理可复用的语言表达单元（Chunk），每个语块包含英文表达、中文翻译和例句。语块掌握度分为三个阶段：activated（刚接触，观看释义和例句）、can_read（能认读，跟读练习）、can_output（能输出，在剧情对话中主动运用）。系统通过usedSceneIds数组追踪语块在多个场景中的使用情况。')
    H(doc, '5.6  AI反馈模块')
    P(doc, 'AI反馈模块接收用户的口语回答文本，拼接包含题目标准答案和行业背景的系统Prompt，调用DeepSeek API的streamText方法，以SSE流式返回四个维度的评分和改进建议。除基础评分外还支持AI教学（Teach）和词汇拓展（WordEnrichment），提供基础/中级/高级三层释义与记忆技巧。评分完成后记录该次反馈的使用记录，用于会员权益控制。')
    H(doc, '5.7  TTS语音模块')
    P(doc, 'TTS模块设计统一的TtsProvider抽象接口，MiniMaxProvider和CartesiaProvider分别实现。合成结果通过配置哈希SHA1（provider + model + voiceId + text + params）缓存至QuestionAudio表，音频文件存储于腾讯云COS，相同配置仅生成一次。支持词级时间戳生成，实现逐词高亮的跟读效果。')
    H(doc, '5.8  互动剧本模块')
    P(doc, '互动剧本模块基于Ink脚本引擎构建交互式剧情学习体验。ScriptModule管理剧本章节（ScriptEpisode），每个章节关联一个场景（Scene）和一组核心词汇/语块。Ink脚本以JSON格式存储，运行时动态解析剧情分支和对话选项。NPC角色与游戏地点关联，用户通过选择对话选项推进剧情。ExplorationModule支持自由对话模式，在游戏地图的不同地点中与NPC对话练习。')
    H(doc, '5.9  模考模块')
    P(doc, '模考模块管理模考试卷（MockPaper）的创建、题目关联（MockPaperQuestion）和考试记录（MockExamRecord）。自动评分功能计算总分和分项得分，根据各题型正确率生成薄弱环节分析，帮助考生针对性强化训练。')
    H(doc, '5.10  会员支付模块')
    P(doc, '会员模块定义套餐信息（MembershipPlan）和用户会员状态（UserMembership）。支付模块创建订单（Order），处理支付网关异步回调，在数据库事务中同步更新订单状态和会员有效期。支持月度/年度两种周期和标准/进阶两级权益。')
    H(doc, '5.11  通知与反馈模块')
    P(doc, '通知模块支持广播通知（面向全体用户）和定向通知（选择特定用户）。通知内容以Markdown格式编辑，可嵌入图片。反馈模块收集用户提交的问题和建议，支持状态标记（待处理/已解決/已关闭）和管理员回复。')
    H(doc, '5.12  管理后台模块')
    P(doc, '管理后台模块提供完整的运营管理功能。用户管理支持查看用户列表、详情和修改角色；会员管理支持查看和取消会员订阅；账单管理展示收入统计和订单明细；题库管理支持对题库、专题、题目的增删改查；资料库管理提供树形文件组织；系统设置通过键值表实现动态配置。所有管理接口通过requireAdmin守卫确保仅管理员可访问。')

    # ═══════ 6 数据库设计 ═══════
    np(doc, pg); pg += 1; H(doc, '6  数据库设计')
    H(doc, '6.1  数据模型总览')
    P(doc, '系统共设计35+个Prisma数据模型，涵盖用户认证、题库内容、场景学习、语块管理、互动剧本、练习行为、模考、会员支付、文件资源、通知反馈、成就排行、等级系统、系统配置等全部业务域。所有模型使用snake_case表名映射，通过@@map注解指定物理表名。')
    H(doc, '6.2  核心表设计')
    T(doc, ['表名','核心字段','约束'],[
        ('user','id, email, name, role','主键，邮箱唯一'),
        ('practice_record','id, userId, questionId, actionType','FK→user, FK→question_item'),
        ('order','id, orderNo, userId, amount, status, paidAt','FK→user, 状态索引'),
        ('question_bank','id, name, province, language, examType','唯一组合约束'),
        ('question_item','id, topicId, title, difficulty, keywords','FK→question_topic, GIN索引'),
        ('scene','id, title, location, requiredLevel','前置场景 + 等级解锁'),
        ('chunk','id, sceneId, english, chinese','场景关联 + usedSceneIds'),
        ('script_episode','id, chapterId, title, inkScript','Ink剧本JSON + 核心语块'),
        ('game_location','id, mapId, name, icon','地图 + NPC关联'),
        ('membership_plan','id, name, level, price, durationDays','枚举level'),
        ('daily_activity','userId, date, count','复合唯一[userId,date]'),
    ])
    H(doc, '6.3  ER关系图')
    diagram(doc, '04-database-er.png', '图3  数据库核心模型ER关系图')
    P(doc, '用户（User）是系统的核心实体，通过外键关联练习记录、模考记录、收藏题目、生词本、订单、会员状态和每日活跃记录。题库（QuestionBank）通过专题（QuestionTopic）关联题目（QuestionItem），题目内容（QuestionContent）以一对一关系存储题目的中英文文本和参考答䅁。游戏化方面，Scene通过前置场景和等级要求形成学习路径，Chunk通过usedSceneIds数组跨场景追踪掌握度，ScriptEpisode通过Ink剧本引擎驱动剧情互动，GameLocation连接NPC角色和地图系统构成探险体验。')
    H(doc, '6.4  索引设计')
    P(doc, '系统采用以下索引策略保障查询性能。所有关联字段（userId、questionId、planId等）自动建立B-tree索引。高频时间范围查询字段（createdAt、paidAt、takenAt）建立降序索引支持最近记录排序。Order表的status字段独立索引支持按支付状态筛选统计。QuestionItem的keywords数组字段使用GIN索引加速标签包含查询。')
    H(doc, '6.5  数据一致性')
    P(doc, '外键约束：所有关联表使用Prisma的@relation + onDelete: Cascade，确保用户删除时级联清理关联子表记录。事务处理：支付关键操作使用Prisma $transaction在数据库事务中原子执行订单创建和会员开通。复合唯一约束：FavoriteQuestion的[userId, questionId]、VocabularyWord的[userId, term]、DailyActivity的[userId, date]等复合约束防止冗余数据。')

    # ═══════ 7 API接口设计 ═══════
    np(doc, pg); pg += 1; H(doc, '7  API接口设计')
    H(doc, '7.1  设计规范')
    P(doc, '接口遵循RESTful设计风格：GET用于查询，POST用于创建，PATCH用于部分更新，DELETE用于删除。业务API统一前缀为/api/v1/guide-exam，认证路由为/api/auth/*。认证采用Bearer Token方案，通过requireAuthSession中间件提取用户身份。URL路径版本化管理，当前版本为v1。')
    H(doc, '7.2  请求处理流程')
    diagram(doc, '05-api-flow.png', '图4  API请求处理时序图')
    P(doc, '客户端请求经Nginx路由分发至NestJS应用。AuthGuard校验Token有效性，ValidationPipe校验请求参数完整性，通过后由Service执行业务逻辑并通过Prisma操作数据库。响应经TransformInterceptor统一封装后返回客户端。异常由AllExceptionsFilter全局捕获并格式化输出。')
    H(doc, '7.3  统一响应格式')
    P(doc, '所有API响应由TransformInterceptor自动封装为统一格式。成功响应格式为{code: 200, message: "Success", data: T}。错误响应格式为{code: 4xx/5xx, message: "错误描述", data: null}。AllExceptionsFilter捕获所有未处理异常，记录请求路径和参数信息后返回格式化错误。ValidationPipe校验失败返回400状态码及具体字段错误描述。')
    H(doc, '7.4  SSE流式响应')
    P(doc, 'AI反馈接口（/practice-ai/feedback）采用Server-Sent Events流式协议。请求流程为：前端POST请求携带题目文本和用户回答，后端拼接包含评分维度的系统Prompt，调用Vercel AI SDK的streamText方法传入DeepSeek模型配置，逐token获取AI响应后通过Response.write()推送SSE格式数据，流结束后写入[DONE]标记并调用end()方法。')

    # ═══════ 8 安全设计 ═══════
    np(doc, pg); pg += 1; H(doc, '8  安全设计')
    H(doc, '8.1  认证流程')
    P(doc, '系统使用Better Auth框架管理用户认证。用户登录成功后生成Session Token存储于HTTP响应中。前端将Token保存至localStorage（键名manyu-bearer-token）。后续每次请求通过Axios请求拦截器自动注入Authorization: Bearer <token>头。后端通过requireAuthSession()函数校验Token有效性并解析用户信息。')
    H(doc, '8.2  权限控制')
    diagram(doc, '05-api-flow.png', '图5 API请求处理流程图')
    P(doc, '系统采用RBAC角色-权限模型。User表中role字段取值为user或admin。普通用户可访问前端业务页面（练习、模考、个人中心等）。管理员额外拥有管理后台全部功能。权限校验通过requireAuthSession和requireAdmin两个守卫函数实现，前者校验会话有效性，后者在校验基础上额外检查角色是否为admin。')
    H(doc, '8.3  数据安全')
    P(doc, '文件上传采用腾讯云COS的STS临时密钥机制。前端通过后端接口获取临时SecretKey后直接上传文件至COS，服务器端永久密钥不对外暴露。文件资产基于SHA256哈希去重，相同内容只存储一份。引用计数管理：FileReference表记录每个文件的引用来源，refCount归零且超过30天的文件由定时任务自动清理。')
    H(doc, '8.4  隐私合规')
    P(doc, '系统提供完整的法律文本页面，包括用户协议、隐私政策、儿童隐私保护政策、第三方SDK列表、信息收集清单、权限说明、ICP备案信息和联系我们，满足个人信息保护法和App违法违规收集使用个人信息行为认定方法的合规要求。用户在个人信息中心可修改账户信息、管理授权、申请注销账户。')

    # ═══════ 9 核心业务流程 ═══════
    np(doc, pg); pg += 1; H(doc, '9  核心业务流程')
    H(doc, '9.1  场景学习流程')
    diagram(doc, '09-scene-learning-flow.png', '图6  场景化学习与探险流程图')
    P(doc, '用户进入场景学习后，首先按类别和等级浏览可用场景。每个场景详情的展示词汇、语块和训练主题。学习路径为：词汇学习（Vocabulary）→ 语块练习（Chunk三阶段）→ 训练主题（TrainingTopic）→ 互动剧本（Ink Script）。词汇和语块掌握后可在NPC对话中主动输出，系统追踪掌握度从activated到can_output的转化。完成场景学习获得经验值，累积提升用户等级。')
    P(doc, '在互动剧本中，用户载入Ink脚本，解析当前剧情节点和对话选项。用户选择对话分支后，系统记录对话记录（ExplorationRecord）并根据剧情推进更新状态。剧本完成后记录成绩（ScriptRecord）。自由对话模式下，用户可在游戏地图的不同地点与NPC直接对话练习。')
    H(doc, '9.2  用户练题流程')
    diagram(doc, '02-business-flow.png', '图7  核心业务流程图')
    P(doc, '用户登录后首先绑定题库，选择省份、语种、考试类型和面试形式，系统加载对应题库内容。题库首页按景点介绍和其他题型分组展示题目。进入练习后支持TTS语音播放、参考答案查看、中文翻译显隐、收藏题目和加入生词等操作。答题完成可请求AI评分反馈，获得多维度的评分和改进建议。')
    H(doc, '9.3  模拟考试流程')
    P(doc, '用户在模考页面查看个人看板数据（平均分、总次数、通过率、最高分）。从试卷列表中选择标准卷或强化卷后进入全屏考试模式。按顺序逐题作答，完成后系统自动计算总分和各题型得分，生成薄弱环节分析报告。历史成绩以趋势折线图展示，帮助用户追踪学习进步情况。')
    H(doc, '9.4  TTS语音流程')
    diagram(doc, '07-tts-flow.png', '图8  TTS语音合成流程图')
    P(doc, '用户点击播放按钮后，系统计算配置哈希值SHA1（provider + model + voiceId + text + params），查询QuestionAudio缓存表。有缓存则直接返回COS音频URL。无缓存则根据当前配置选择MiniMax或Cartesia厂商，调用TTS API生成音频，上传至COS同时写入缓存记录。整个流程异步执行，不阻塞用户界面操作。')
    H(doc, '9.5  AI反馈时序')
    diagram(doc, '05-api-flow.png', '图9  AI口语反馈时序图')
    P(doc, '用户点击AI反馈按钮后，前端发起POST请求携带题目文本和用户回答。后端拼接系统Prompt后调用DeepSeek API的streamText方法。AI分析结果以SSE流式方式逐token推送至前端，前端通过ReadableStream实时接收并逐字渲染评分文本。评分完成后后端记录使用情况用于会员权益控制。')
    H(doc, '9.6  支付开通流程')
    diagram(doc, '08-payment-flow.png', '图10  支付与会员开通时序图')
    P(doc, '用户选择套餐后点击开通，后端创建Order记录并返回支付二维码。用户扫码支付后支付网关异步回调通知后端。后端在事务中更新订单状态为已支付，同时创建或延长用户会员记录。会员权益即时生效，用户可在个人中心查看到期时间。若支付失败，返回失败原因供用户再次尝试。')

    # ═══════ 10 部署与运维 ═══════
    np(doc, pg); pg += 1; H(doc, '10  部署与运维')
    H(doc, '10.1  CI/CD流程')
    P(doc, '代码推送至GitHub main分支后自动触发GitHub Actions工作流。流程依次执行依赖安装（pnpm install）、代码检查（lint + typecheck）、多阶段Docker构建、镜像推送至Docker Registry、SSH连接云服务器拉取最新镜像并重启服务。整个流程约3-5分钟完成，确保代码变更安全快速地部署至生产环境。')
    H(doc, '10.2  容器化方案')
    P(doc, '后端使用多阶段Dockerfile构建。第一阶段安装全量依赖并编译TypeScript源码。第二阶段仅复制编译产物和生产依赖，基于node:22-alpine精简基础镜像，最终体积约200MB。前端同样采用多阶段构建，编译后的静态资源通过独立Nginx容器提供服务，镜像体积约50MB。开发环境使用Docker Compose编排多容器服务。')
    H(doc, '10.3  部署架构')
    diagram(doc, '06-deployment.png', '图11  系统部署架构图')
    P(doc, '生产环境使用Docker Compose编排三个容器：Nginx负责反向代理和前端静态资源服务，NestJS Backend提供API服务，PostgreSQL存储关系数据。Nginx配置/api路径反向代理至Backend容器，其他路径直接返回前端index.html（SPA路由支持）。音频和图片等静态文件通过腾讯云COS的CDN加速分发。')

    # ═══════ 11 技术难点与解决方案 ═══════
    np(doc, pg); pg += 1; H(doc, '11  技术难点与解决方案')
    H(doc, '11.1  TTS多厂商抽象')
    P(doc, '难点：不同TTS厂商的API接口差异较大（MiniMax返回PCM裸流需要转码、Cartesia支持WebSocket流式返回WAV），需要实现统一的调用方式并避免重复调用产生费用浪费。')
    P(doc, '解决方案：定义TtsProvider抽象接口，包含synthesize(text, voice, options)方法和getVoices()方法。各厂商分别实现该接口，通过配置或策略模式动态选择厂商。配置哈希缓存机制：hash = SHA1(provider + model + voiceId + text + JSON.stringify(params))，QuestionAudio表以[questionId, configHash]为唯一键，命中即返回，完全消除重复调用。')
    H(doc, '11.2  AI流式反馈')
    P(doc, '难点：大语言模型推理时间较长（5-15秒），同步请求模式下用户需等待完整响应，体验割裂。')
    P(doc, '解决方案：采用SSE（Server-Sent Events）协议实现流式传输。后端使用Vercel AI SDK的streamText函数封装对DeepSeek API的流式调用，结果以AsyncIterable形式逐token输出。通过NestJS Response的write()方法以text/event-stream格式逐字推送。前端fetch API通过response.body.getReader()读取ReadableStream，每收到data:前缀事件后立即解析渲染到页面。')
    H(doc, '11.3  Ink剧本引擎')
    P(doc, '难点：互动剧情需要动态解析剧本分支逻辑，实现不同选择导致不同剧情走向。传统硬编码方式难以维护和扩展剧情内容。')
    P(doc, '解决方案：采用Ink脚本语言作为剧本描述格式。Ink脚本编译为JSON格式存储于数据库（InkScript模型），运行时通过解析引擎动态读取剧情节点和选项分支。ScriptEpisode关联InkScript和场景，核心词汇/语块在剧情中自然呈现。NPC角色与游戏地点通过GameLocation和LocationNpc关联，支持自由对话和剧情驱动两种模式。')
    H(doc, '11.4  Monorepo管理')
    P(doc, '难点：前后端代码在同一个仓库中维护但运行时依赖各不相同（后端需要NestJS、Prisma，前端需要React、Vite），构建时需要隔离依赖并共享类型定义。')
    P(doc, '解决方案：使用pnpm workspace定义@manyu/backend和@manyu/frontend两个工作区包，各自维护独立的package.json。类型共享通过直接引用源文件路径实现。Docker构建时使用pnpm --filter精确安装目标子项目依赖，避免将前端node_modules打包到后端镜像。开发时使用pnpm --parallel -r dev同时启动前后端开发服务器。')
    H(doc, '11.5  语块三阶段掌握度模型')
    P(doc, '难点：传统的掌握度模型（如0-100分）缺乏可操作的学习阶段指引，学习者不清楚自己当前应该处于什么状态、下一步该做什么。')
    P(doc, '解决方案：设计三阶段语块掌握度模型：activated（刚接触，已查看语块释义和例句）→ can_read（能认读，已跟读练习）→ can_output（能输出，已在剧情对话中主动运用）。每个阶段对应明确的学习行为指标。usedSceneIds数组记录语块被使用的场景ID列表，支持跨场景的掌握度评估。系统在剧情对话训练中自动检查并推进用户掌握度阶段。')

    H(doc, '11.6  文件去重')
    P(doc, '难点：同一音频文件可能被多个题目引用，用户上传的图片可能被多处使用。简单的引用计数不做去重会导致存储浪费和关联混乱。')
    P(doc, '解决方案：FileAsset表以sha256字段为唯一键，相同哈希值只存储一个物理文件。FileReference表记录每条引用关系，refCount为对应assetId的引用计数。定时任务每分钟扫描FileAsset表，找出refCount=0且创建超过30分钟的文件，标记过期后删除COS文件并清理数据库记录。')

    H(doc, '11.7  跨端兼容')
    P(doc, '难点：系统需同时支持Web浏览器和iOS原生应用，路由模式、音频播放、本地存储等基础能力在不同平台存在差异。')
    P(doc, '解决方案：路由使用HashRouter，URL格式#/path在静态部署和Capacitor iOS下表现一致。音频播放设计PlaybackAdapter抽象接口，Web端实现基于HTMLAudioElement，iOS端使用Capacitor插件。存储使用统一异步Storage接口，Web端基于localStorage，iOS端使用Capacitor Preferences插件。')

    # ── Save ──
    doc.save(OUTPUT)
    print(f'  Done: {OUTPUT}')

if __name__ == '__main__':
    main()
