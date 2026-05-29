"""
漫语町（ManYu）— 软件著作权申请文档生成器
生成符合中国版权保护中心要求的 Word 文档
"""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import datetime

doc = Document()

# ─── 全局样式 ───────────────────────────────────────────────
style = doc.styles['Normal']
font = style.font
font.name = 'SimSun'
font.size = Pt(12)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.18)
    section.right_margin = Cm(3.18)

def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)
    return h

def add_para(text, bold=False, align=None, size=None, space_after=Pt(6)):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    if size:
        run.font.size = size
    if align:
        p.alignment = align
    p.paragraph_format.space_after = space_after
    return p

def add_table_row(table, cells, bold=False):
    row = table.add_row()
    for i, cell_text in enumerate(cells):
        row.cells[i].text = ''
        p = row.cells[i].paragraphs[0]
        run = p.add_run(cell_text)
        run.bold = bold
        run.font.size = Pt(11)
        p.paragraph_format.space_after = Pt(2)

# ══════════════════════════════════════════════════════════════
# 封面
# ══════════════════════════════════════════════════════════════

doc.add_paragraph()
doc.add_paragraph()
add_para('计算机软件著作权登记申请', bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=Pt(22))
doc.add_paragraph()
add_para('软件名称：漫语町（ManYu）沉浸式英语输出训练系统', bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=Pt(14))
add_para('版本号：V1.0', align=WD_ALIGN_PARAGRAPH.CENTER, size=Pt(12))
add_para(f'编制日期：{datetime.date.today().strftime("%Y年%m月%d日")}', align=WD_ALIGN_PARAGRAPH.CENTER, size=Pt(12))

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 一、软件基本信息
# ══════════════════════════════════════════════════════════════

add_heading('一、软件基本信息', level=1)

table = doc.add_table(rows=0, cols=2)
table.style = 'Light Grid Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.columns[0].width = Cm(4.5)
table.columns[1].width = Cm(11)

info = [
    ('软件全称', '漫语町（ManYu）沉浸式英语输出训练系统'),
    ('软件简称', '漫语町 / ManYu'),
    ('版本号', 'V1.0'),
    ('软件分类', '应用软件 — 教育学习类'),
    ('开发完成日期', '2026年5月27日'),
    ('首次发表日期', '未发表'),
    ('开发方式', '独立开发'),
    ('运行环境', 'Web浏览器 + iOS 原生（Capacitor）'),
    ('操作系统', '跨平台（Windows / macOS / iOS / Android）'),
    ('编程语言', 'TypeScript / JavaScript'),
    ('源程序量', '约 80,000 行'),
]
for key, val in info:
    add_table_row(table, [key, val])

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════
# 二、开发者信息
# ══════════════════════════════════════════════════════════════

add_heading('二、开发者信息', level=1)
add_para('申请人：________（请填写个人姓名或公司名称）')
add_para('联系地址：________')
add_para('联系电话：________')
add_para('邮箱：________')

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════
# 三、软件开发目的
# ══════════════════════════════════════════════════════════════

add_heading('三、软件开发目的', level=1)
add_para(
    '漫语町（ManYu）是一款沉浸式英语输出训练系统。'
    '以"场景驱动、语块积累、互动输出"为核心学习理念，将英语真实交流场景融入练习流程，'
    '帮助学习者在沉浸式互动中高效提升英语口语输出能力。随着全球化进程的加速，'
    '英语实用口语能力的需求日益旺盛。然而，传统学习方式面临以下痛点：'
)
bullets = [
    '缺乏真实英语交流场景的沉浸式训练环境',
    '口语练习缺乏即时纠错和升级反馈',
    '词汇和表达积累碎片化，缺乏系统性组织',
    '学习进度与效果缺乏量化追踪，学习动力难以持续',
]
for b in bullets:
    doc.add_paragraph(b, style='List Bullet')

add_para(
    '基于上述背景，本系统旨在通过数字化手段和游戏化设计，为英语学习者提供一站式的沉浸式英语输出训练解决方案。'
    '系统以场景化学习为核心，通过词汇积累、语块训练、AI纠错反馈、'
    'TTS语音合成、互动剧本探险等核心功能，帮助学习者在真实语境中高效提升英语口语输出能力。'
)

# ══════════════════════════════════════════════════════════════
# 四、软件主要功能
# ══════════════════════════════════════════════════════════════

add_heading('四、软件主要功能', level=1)

features = [
    ('1. 题库浏览与管理',
     '按省份、语种、考试类型、面试形式等多维度筛选题库，支持关键词搜索。'
     '题库内容涵盖情景对话、话题陈述、日常表达、学术讨论等学习类别，'
     '每道题目包含中英文题目、参考答案、翻译、关键词、重点词汇等结构化信息。'),
    ('2. 场景化学习',
     '以"场景（Scene）"为学习单元，每个场景包含场景词汇（Vocabulary）、语块（Chunk）和训练主题（TrainingTopic）。'
     '语块学习采用三阶段掌握度模型：激活（activated）→ 认读（can_read）→ 输出（can_output），'
     '配合例句、翻译和语义上下文帮助深度理解。'),
    ('3. 互动剧本与探险',
     '基于 Ink 脚本引擎的互动式剧情学习，在不同场景地图（GameMap）和地点（Location）中与 NPC 角色展开对话练习。'
     '支持自由对话模式和剧情驱动模式，对话记录自动保存并可回溯复习。'
     '每个剧本章节（ScriptEpisode）设有核心词汇和语块，在剧情互动中自然习得。'),
    ('4. 口语练习',
     '逐题练习模式，支持学习模式（默认显示答案）和练习模式（按题切换显示答案）两种训练方式。'
     '内置音频播放（浏览器TTS）、单词点击释义、中文翻译显隐切换等功能。'
     '支持快捷键操作（播放、答案、翻译、收藏、上下题），提高练习效率。'),
    ('5. AI 口语评分与教学',
     '集成 DeepSeek 大语言模型，对用户的口语回答进行实时评分（发音、语法、内容完整性），'
     '并给出针对性的改进建议。同时支持 AI 教学和词汇拓展（Word Enrichment），'
     '提供基础/中级/高级三层释义与记忆技巧。采用 SSE 流式响应技术，反馈结果逐字输出。'),
    ('6. TTS 语音合成',
     '集成 MiniMax 和 Cartesia 两大 TTS 厂商，支持多语种、多音色的语音合成。'
     '系统自动缓存已生成的音频文件（按配置哈希去重），避免重复请求，提升加载速度。'
     '支持词级时间戳，实现逐词高亮的跟读体验。'),
    ('7. 模拟考试',
     '标准化试卷模考，支持标准卷和强化卷两种难度。'
     '考试结束后自动评分，记录各题型得分与薄弱环节分析，帮助考生针对性强化。'),
    ('8. 学习档案与数据看板',
     '个人学习中心：练习热力图、掌握度进度条、错题本、收藏夹、生词本。'
     '数据统计：总练习量、连续打卡天数、模考历史成绩、掌握度分布。'
     '游戏化元素：经验值（Experience）与等级系统、成就徽章系统、学习排行榜。'),
    ('9. 会员系统',
     '支持标准会员和进阶会员两级套餐，按月/年订阅。'
     '会员权益包括完整题库访问、不限次 AI 点评、录音上传、错题分析、优先客服等。'),
    ('10. 管理后台',
     'Web 管理后台支持用户管理、会员管理、账单管理、消息通知、题库内容管理、'
     '资料库管理、反馈管理、数据统计看板、系统设置等功能，方便运营人员维护系统内容与配置。'),
]

for title, desc in features:
    add_para(title, bold=True)
    add_para(desc)

# ══════════════════════════════════════════════════════════════
# 五、技术特点
# ══════════════════════════════════════════════════════════════

add_heading('五、技术特点', level=1)

add_para('1. 现代化全栈技术架构', bold=True)
add_para(
    '采用 Monorepo（pnpm workspace）组织前后端代码，'
     '后端基于 NestJS + Prisma + PostgreSQL 构建健壮的 API 服务层，共 25 个业务模块；'
    '前端基于 React 19 + TypeScript + Vite，搭配 shadcn/ui 组件库和 Tailwind CSS，'
    '实现美观、可访问的用户界面。前后端共享类型定义，确保接口一致性。'
)

add_para('2. 先进的认证系统', bold=True)
add_para(
     '采用 Better Auth 认证框架，支持邮箱密码注册登录、OTP 邮箱验证、'
    '会话管理、角色权限控制（用户/管理员）。所有业务接口通过 requireAuthSession 中间件统一鉴权。'
)

add_para('3. 多厂商 TTS 引擎抽象', bold=True)
add_para(
    '设计统一的 TTS Service 抽象层，支持 MiniMax、Cartesia 等多厂商无缝切换。'
    '基于配置哈希（SHA1）的自动音频缓存机制，同一文本+配置组合仅请求一次，'
    '显著降低 TTS API 调用成本。采用引用计数的文件资产管理，自动清理过期文件。'
)

add_para('4. Ink 互动剧本引擎', bold=True)
add_para(
    '集成 Ink 脚本语言解析引擎，用于创作交互式对话剧本。剧本以 JSON 格式存储 Ink 脚本内容，'
    '运行时动态解析剧情分支和对话选项。NPC 角色与多个游戏地点关联，用户通过选择对话选项推进剧情，'
    '在沉浸式故事体验中自然习得英语场景表达。'
)

add_para('5. 语块三阶段掌握度模型', bold=True)
add_para(
    '创新性地引入语块（Chunk）掌握度三阶段模型：activated（刚接触）→ can_read（能认读）→ can_output（能输出）。'
    '每个阶段对应不同的学习行为：激活阶段观看语块释义和例句；认读阶段跟读练习；输出阶段在剧情对话中主动运用。'
    '系统追踪每个语块在各场景中的使用次数（usedSceneIds），实现跨场景的掌握度评估。'
)

add_para('6. 游戏化学习体系', bold=True)
add_para(
    '构建完整的游戏化学习闭环：用户通过场景学习获得经验值（Experience），经验值累积提升用户等级（Level）。'
    '每个场景设置前置解锁条件（prerequisiteScenes + 等级要求），引导学习者循序渐进。'
    '成就系统（Achievement）自动检测练习天数、题目数量、模考成绩等里程碑并解锁徽章。'
    '排行榜（Leaderboard）从练习量、模考分、连续打卡三个维度激励学习。'
)

add_para('7. AI 流式反馈', bold=True)
add_para(
    '使用 Vercel AI SDK 的 streamText 能力，结合 DeepSeek 大模型，'
    '实现口语评分的流式输出（SSE）。用户无需等待完整分析结果，'
    '即可逐字看到评分反馈，体验媲美实时对话。'
)

add_para('8. 响应式与跨端支持', bold=True)
add_para(
    '前端基于 React + Tailwind CSS 实现完全响应式设计，桌面端和移动端自适应布局。'
    '通过 Capacitor 框架打包为 iOS 原生应用，支持音频录制、本地存储等原生能力。'
    '采用 Hash 路由模式，兼容静态部署和 iOS 平台。'
)

add_para('9. 数据安全与合规', bold=True)
add_para(
    '用户密码经 Better Auth 自动加盐哈希；API 请求通过 Bearer Token 鉴权；'
    '支持腾讯云 COS 临时密钥（STS）的前端直传，避免密钥泄露；'
    '文件资产基于 SHA256 去重，引用计数管理生命周期。'
    '提供用户协议、隐私政策、儿童隐私保护等法律文本。'
)

add_para('10. 灵活的数据库设计', bold=True)
add_para(
    '使用 Prisma ORM 管理 PostgreSQL 数据库，35+ 个数据模型覆盖用户、题库、练习、模考、'
    '场景、语块、剧本、探索、会员、支付、通知、文件资产、成就、排行榜、等级等完整业务域。'
    '支持复合唯一约束防止重复数据，采用 Cascade 删除策略维护数据一致性。'
    'SystemConfig 键值表支持动态配置，无需改表即可新增配置项。'
)

add_para('11. 自动化部署与运维', bold=True)
add_para(
    '基于 Docker 多阶段构建（NestJS + Nginx）优化生产镜像体积。'
    'GitHub Actions CI/CD 自动化测试、构建与部署流程。'
    'Nginx 反向代理统一入口，前端静态资源 CDN 加速，后端 API 负载均衡。'
)

# ══════════════════════════════════════════════════════════════
# 六、软件运行环境
# ══════════════════════════════════════════════════════════════

add_heading('六、软件运行环境', level=1)

env_data = [
    ('服务端', 'Node.js >= 22, PostgreSQL >= 16'),
    ('客户端', 'Chrome / Safari / Edge 等现代浏览器'),
    ('移动端', 'iOS 16+ (Capacitor)'),
    ('开发工具', 'VS Code, pnpm >= 9'),
    ('SDK/框架', 'NestJS 11, React 19, Prisma 6, Tailwind CSS 4, Vite 6'),
    ('外部服务', 'DeepSeek API, MiniMax TTS, Cartesia TTS, 腾讯云 COS'),
]

env_table = doc.add_table(rows=0, cols=2)
env_table.style = 'Light Grid Accent 1'
env_table.alignment = WD_TABLE_ALIGNMENT.CENTER
env_table.columns[0].width = Cm(4)
env_table.columns[1].width = Cm(11.5)

for key, val in env_data:
    add_table_row(env_table, [key, val])

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════
# 七、软件结构图
# ══════════════════════════════════════════════════════════════

add_heading('七、软件架构与目录结构', level=1)
add_para(
    '系统采用前后端分离的 Monorepo 架构，代码按功能域模块化组织。'
    '主要目录结构如下：'
)

code_text = """manyu/
├── apps/
│   ├── backend/           # NestJS 后端服务
│   │   ├── prisma/        # 数据库 Schema + Migrations
│   │   └── src/
│   │       ├── common/    # 公共层（拦截器、过滤器、Prisma）
│   │       └── modules/   # 业务模块（auth, practice, mock-exam, tts...）
│   └── frontend/          # React 前端
│       ├── capacitor/     # iOS 原生项目 (Capacitor)
│       └── src/
│           ├── components/# UI 组件（shadcn/ui）
│           ├── features/  # 按功能域组织（api + pages）
│           ├── stores/    # Zustand 状态管理
│           └── providers/ # Context Provider
├── docker/                # Dockerfile + Nginx 配置
├── docs/                  # 项目文档
└── docker-compose.yml     # 本地开发编排"""

p = add_para('', space_after=Pt(6))
run = p.add_run(code_text)
run.font.name = 'Consolas'
run.font.size = Pt(9)

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════
# 八、源程序代码说明
# ══════════════════════════════════════════════════════════════

add_heading('八、源程序代码', level=1)
add_para(
    '按照中国版权保护中心要求，本申请附带的源程序代码为：'
)
bullets2 = [
    '前、后各连续 30 页源程序代码（共 60 页），每页不少于 50 行',
    '代码取自本项目核心模块，包括：后端 API 服务（NestJS Controller/Service）、'
    'Prisma 数据库模型定义、前端核心页面组件等',
    '代码以 A4 纸单面打印，页眉标注软件名称及版本号，页码连续编号',
    '源程序代码的页眉格式为："漫语町（ManYu）沉浸式英语输出训练系统 V1.0  源代码  第 X 页"',
]
for b in bullets2:
    doc.add_paragraph(b, style='List Bullet')

# ══════════════════════════════════════════════════════════════
# 九、用户操作手册说明
# ══════════════════════════════════════════════════════════════

add_heading('九、用户操作手册', level=1)
add_para(
    '按照中国版权保护中心要求，本申请附带的用户操作手册为：'
)
bullets3 = [
    '前、后各连续 30 页用户操作手册（共 60 页）',
    '内容包括：系统概述、运行环境说明、用户注册与登录、题库管理与练习、'
    '模拟考试操作、AI 反馈使用说明、个人中心功能介绍、管理后台操作指南等',
    '手册以 A4 纸单面打印，页眉标注软件名称及版本号，页码连续编号',
    '操作手册的页眉格式为："漫语町（ManYu）沉浸式英语输出训练系统 V1.0  操作手册  第 X 页"',
]
for b in bullets3:
    doc.add_paragraph(b, style='List Bullet')

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════
# 十、签署页
# ══════════════════════════════════════════════════════════════

add_heading('十、申请人声明', level=1)
add_para(
    '本人（单位）郑重声明：'
    '所提交的登记申请材料及填写内容真实、准确、完整，'
    '并对所填报的内容承担相应法律责任。'
)
doc.add_paragraph()
doc.add_paragraph()

add_para('申请人（签章）：____________________')
add_para(f'日期：{datetime.date.today().strftime("%Y年%m月%d日")}')

# ══════════════════════════════════════════════════════════════
# 保存
# ══════════════════════════════════════════════════════════════

output_path = r'C:\Users\z1309\Desktop\work\speakGuild\docs\软件著作权\ManYu软件著作权申请文档.docx'
doc.save(output_path)
print(f'✅ 已生成：{output_path}')
