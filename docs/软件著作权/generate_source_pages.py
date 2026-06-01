"""
漫语町（ManYu）软件著作权 — 源程序代码生成器
从项目实际源文件中提取代码，剔除注释，生成符合版权中心要求的纯代码 60 页文档
"""

import os
import re
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

PROJECT_ROOT = r'C:\Users\z1309\Desktop\work\speakGuild'
OUTPUT_PATH = r'C:\Users\z1309\Desktop\work\speakGuild\docs\软件著作权\ManYu源代码.docx'
HEADER_TEXT = '漫语町（ManYu）沉浸式英语输出训练系统 V1.0  源代码'
SOFTWARE_NAME = '漫语町（ManYu）沉浸式英语输出训练系统 V1.0'
LINES_PER_PAGE = 55
TOTAL_PAGES = 60

SOURCE_FILES = [
    'apps/backend/src/main.ts',
    'apps/backend/src/app.module.ts',
    'apps/backend/src/common/prisma/prisma.service.ts',
    'apps/backend/src/common/interceptors/transform.interceptor.ts',
    'apps/backend/src/common/filters/all-exceptions.filter.ts',
    'apps/backend/src/common/response/api-response.ts',
    'apps/backend/prisma/schema.prisma',
    # Auth
    'apps/backend/src/modules/auth/auth.module.ts',
    'apps/backend/src/modules/auth/session.util.ts',
    'apps/backend/src/modules/auth/password.service.ts',
    # Admin
    'apps/backend/src/modules/admin/admin.controller.ts',
    'apps/backend/src/modules/admin/admin.service.ts',
    'apps/backend/src/modules/admin/admin-stats.service.ts',
    'apps/backend/src/modules/admin/system-config/system-config.service.ts',
    # Practice & Mock
    'apps/backend/src/modules/practice/practice.service.ts',
    'apps/backend/src/modules/practice-ai/practice-ai.service.ts',
    'apps/backend/src/modules/mock-exam/mock-exam.service.ts',
    # TTS
    'apps/backend/src/modules/tts/tts.service.ts',
    # Scene & Chunk & Script (New)
    'apps/backend/src/modules/scene/scene.service.ts',
    'apps/backend/src/modules/chunk/chunk.service.ts',
    'apps/backend/src/modules/script/script.service.ts',
    'apps/backend/src/modules/exploration/exploration.service.ts',
    'apps/backend/src/modules/learning/learning.service.ts',
    'apps/backend/src/modules/level/level.service.ts',
    'apps/backend/src/modules/onboarding/onboarding.service.ts',
    # Membership & Pay
    'apps/backend/src/modules/membership/membership.service.ts',
    'apps/backend/src/modules/pay/pay.service.ts',
    'apps/backend/src/modules/coupon/coupon.service.ts',
    'apps/backend/src/modules/referral/referral.service.ts',
    # Notifications & Feedback
    'apps/backend/src/modules/notification/notification.service.ts',
    'apps/backend/src/modules/feedback/feedback.service.ts',
    'apps/backend/src/modules/achievement/achievement.service.ts',
    'apps/backend/src/modules/leaderboard/leaderboard.service.ts',
    'apps/backend/src/modules/file-assets/file-assets.service.ts',
    'apps/backend/src/modules/resource-library/resource-library.service.ts',
    # Frontend
    'apps/frontend/src/App.tsx',
    'apps/frontend/src/main.tsx',
    'apps/frontend/src/lib/request.ts',
    'apps/frontend/src/lib/cn.ts',
    'apps/frontend/src/stores/config.store.ts',
    'apps/frontend/src/stores/assets.store.ts',
    'apps/frontend/src/stores/preferences.store.ts',
    'apps/frontend/src/stores/layout.store.ts',
    'apps/frontend/src/providers/auth-provider.tsx',
    'apps/frontend/src/providers/auth-route-guard.tsx',
    'apps/frontend/src/providers/theme-provider.tsx',
    'apps/frontend/src/layout/root-layout.tsx',
    'apps/frontend/src/layout/admin-layout.tsx',
    'apps/frontend/src/features/admin/pages/admin-analytics-page.tsx',
    'apps/frontend/src/features/admin/pages/admin-settings-page.tsx',
    'apps/frontend/src/features/admin/api.ts',
    'apps/frontend/src/features/scene/pages/scene-list-page.tsx',
    'apps/frontend/src/features/practice/pages/practice-page.tsx',
    'apps/frontend/src/features/scene/pages/scene-list-page.tsx',
    'apps/frontend/src/features/profile/pages/profile-page.tsx',
    'apps/frontend/src/features/membership/pages/member-page.tsx',
    'apps/frontend/src/features/mock-exam/pages/mock-page.tsx',
    'apps/frontend/src/components/ui/switch.tsx',
    'apps/frontend/src/components/ui/tabs.tsx',
    'apps/frontend/src/components/ui/card.tsx',
]


def strip_comments(line):
    """Remove comments from a source line, keeping the code part."""
    line = line.rstrip('\n').rstrip('\r')
    stripped = line.strip()
    if stripped.startswith('//') or stripped.startswith('#'):
        return ''
    if stripped.startswith('/*') or stripped.startswith('*'):
        return ''
    # inline comment removal: // within a line
    in_string = False
    string_char = None
    for i, ch in enumerate(line):
        if ch in ('"', "'", '`') and (i == 0 or line[i-1] != '\\'):
            if not in_string:
                in_string = True
                string_char = ch
            elif ch == string_char:
                in_string = False
        if not in_string and i < len(line) - 1 and line[i:i+2] == '//':
            result = line[:i].rstrip()
            return result if result else ''
    return line if line.strip() else ''


def add_page_header(section, page_num):
    header = section.header
    header.is_linked_to_previous = False
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'{HEADER_TEXT}  第 {page_num} 页')
    run.font.name = 'SimSun'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(128, 128, 128)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')


def add_page_footer(section):
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(SOFTWARE_NAME)
    run.font.name = 'SimSun'
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(128, 128, 128)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')


def set_page_margins(section):
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)


def add_code_page(doc, code_lines, page_num):
    section = doc.add_section()
    set_page_margins(section)
    add_page_header(section, page_num)
    add_page_footer(section)

    for line in code_lines:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = Pt(11)
        run = p.add_run(line)
        run.font.name = 'Consolas'
        run.font.size = Pt(7.5)


def read_source_file(filepath):
    full_path = os.path.join(PROJECT_ROOT, filepath)
    if not os.path.exists(full_path):
        return []
    try:
        with open(full_path, 'r', encoding='utf-8') as f:
            return f.readlines()
    except Exception:
        return []


def main():
    print('  Generating source code document...')
    doc = Document()

    # No title page — just pure code

    # ─── Collect all code lines (stripped of comments) ───
    all_lines = []
    for filepath in SOURCE_FILES:
        lines = read_source_file(filepath)
        if not lines:
            continue
        for line in lines:
            cleaned = strip_comments(line)
            if cleaned:
                all_lines.append(cleaned)

    total = len(all_lines)
    pages_needed = min((total + LINES_PER_PAGE - 1) // LINES_PER_PAGE, TOTAL_PAGES)
    print(f'  {total} comment-free lines, generating {pages_needed} pages')

    for pg in range(pages_needed):
        start = pg * LINES_PER_PAGE
        page_lines = all_lines[start:start + LINES_PER_PAGE]
        if not page_lines:
            break
        add_code_page(doc, page_lines, pg + 1)

    doc.save(OUTPUT_PATH)
    print(f'  Done: {OUTPUT_PATH} ({pages_needed} pages)')


if __name__ == '__main__':
    main()
