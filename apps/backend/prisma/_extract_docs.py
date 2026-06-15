import json, sys, os

out_path = r'c:\Users\z1309\Desktop\work\speakGuild\apps\backend\prisma\_extract_out.txt'
results = []

# File 1: 600 common phrases
f1 = r'c:\Users\z1309\Documents\xwechat_files\wxid_5pg1904cumf012_bc58\msg\file\2026-06\600常见词组、翻译、例句（邵艾伦原创）.docx'
results.append(f'=== File 1: 600 common phrases ===')

try:
    from docx import Document
    doc = Document(f1)
    results.append(f'Paragraphs: {len(doc.paragraphs)}')
    for i, p in enumerate(doc.paragraphs):
        t = p.text.strip()
        if t and i < 300:
            results.append(f'{i}: {t[:200]}')
except Exception as e:
    results.append(f'Error: {e}')

results.append('')

# File 2: 400 fixed collocations
f2 = r'c:\Users\z1309\Documents\xwechat_files\wxid_5pg1904cumf012_bc58\msg\file\2026-06\400固定搭配、翻译、例句（邵艾伦原创）.docx'
results.append(f'=== File 2: 400 fixed collocations ===')

try:
    doc2 = Document(f2)
    results.append(f'Paragraphs: {len(doc2.paragraphs)}')
    for i, p in enumerate(doc2.paragraphs):
        t = p.text.strip()
        if t and i < 300:
            results.append(f'{i}: {t[:200]}')
except Exception as e:
    results.append(f'Error: {e}')

with open(out_path, 'w', encoding='utf-8') as f:
    f.write('\n'.join(results))
print('Done', file=sys.stderr)
